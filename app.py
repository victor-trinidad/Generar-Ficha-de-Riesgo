import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm # Importar Cm para mayor precisión
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import io 
import os 

# --- CONFIGURACIÓN ESPECÍFICA DEL ARCHIVO ---
ARCHIVO_EXCEL = 'LMM_ORG_04 Rev. 00 - Matriz Institucional de Gestión de Riesgos.xlsx'
NOMBRE_HOJA = 'LMM_ORG_04' 
FILA_ENCABEZADOS = 16 

# Renombrar columnas para facilitar el acceso
COLUMNAS_MAP = [
    'Num_Riesgo', 'Entorno_Control', 'Origen_Area', 'Proceso_Documento', 
    'Riesgo_Identificado', 'Impacto_Potencial', 'Efecto', 
    'Gravedad', 'Probabilidad', 'PxG', 'Escala_Riesgo', 
    'Control_Existente', 'Tipo_Control', 'Responsable_Seguimiento', 
    'Eficacia_Seguimiento', 'Version', 'Estado_Control', 
    'Acciones', 'Fecha_Identificacion', 'Ultima_Revision'
]

# --- FUNCIONES AUXILIARES DE GENERACIÓN ---

def agregar_seccion_tabla(document, titulo, datos_dict):
    """Agrega una sección formal usando una tabla de dos columnas (Título del Campo | Valor del Campo)."""
    document.add_heading(titulo, level=2)
    
    # Crear una tabla de 2 columnas
    tabla = document.add_table(rows=len(datos_dict), cols=2)
    tabla.style = 'Table Grid'
    
    # Establecer ancho para el título (Columna 0)
    tabla.columns[0].width = Inches(2)
    
    i = 0
    for key, value in datos_dict.items():
        row_cells = tabla.rows[i].cells
        
        # Título del Campo (negrita)
        run_key = row_cells[0].paragraphs[0].add_run(f'{key}')
        run_key.bold = True
        run_key.font.name = 'Arial'
        
        # Valor del Campo (texto normal)
        run_value = row_cells[1].paragraphs[0].add_run(str(value))
        run_value.font.name = 'Arial'
        i += 1
    document.add_paragraph()

def generar_ficha_docx(datos_riesgo):
    """
    Genera la ficha A4 con ajustes precisos de márgenes, encabezado y pie de página.
    """
    document = Document()
    
    # --------------------------------------------------------
    # 1. AJUSTE DE ESTILOS BASE Y PÁGINA (A4)
    # --------------------------------------------------------
    section = document.sections[0]
    
    # Márgenes de Disposición
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Posición de Encabezado/Pie
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.6)

    # Establecer estilo de fuente base (Arial 10pt) para el 'Normal' Style
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    # --------------------------------------------------------
    # 2. ENCABEZADO (TABLA COMPLEJA)
    # --------------------------------------------------------
    header = section.header
    
    # Usaremos una tabla de 4 filas x 4 columnas para replicar el formato
    header_table = header.add_table(4, 4)
    header_table.style = 'Table Grid'
    
    # Anchos de Columna (Usando Cm para exactitud)
    header_table.columns[0].width = Cm(4.3) 
    header_table.columns[1].width = Cm(8.18) 
    header_table.columns[2].width = Cm(1.8) 
    header_table.columns[3].width = Cm(2.7) 

    # Alto de Fila (Intentar 0.5 cm para las 4 filas, que suman 2.0 cm, cercano al 2.05 cm total)
    for row in header_table.rows:
        row.height = Cm(0.5)

    # ------------------------------------------------------------------
    # A. MERGE Y CONTENIDO: Columna 1 (Logo) - Fusiones R0 a R3
    # ------------------------------------------------------------------
    a1 = header_table.cell(0, 0)
    d1 = header_table.cell(3, 0)
    a1.merge(d1)
    
    p_logo = a1.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        # Ajustar tamaño del logo para que quepa en 4.3cm
        p_logo.add_run().add_picture('logo.png', width=Cm(3.0)) 
    except FileNotFoundError:
        p_logo.add_run("Lqf").bold = True
        
    # ------------------------------------------------------------------
    # B. MERGE Y CONTENIDO: Columna 2 (Título/Info) - Fusiones R0 a R1 y R2 a R3
    # ------------------------------------------------------------------
    
    # R0 + R1 (Título Principal LISTADO MAESTRO O MATRIZ)
    b1 = header_table.cell(0, 1)
    b2 = header_table.cell(1, 1)
    b1.merge(b2)
    p_title_top = b1.paragraphs[0]
    p_title_top.add_run('LISTADO MAESTRO O MATRIZ').bold = True
    p_title_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title_top.style.font.size = Pt(10)
    p_title_top.style.font.name = 'Arial'
    
    # R2 + R3 (Título Secundario)
    b3 = header_table.cell(2, 1)
    b4 = header_table.cell(3, 1)
    b3.merge(b4)
    p_title_bottom = b3.paragraphs[0]
    p_title_bottom.add_run("MATRIZ INSTITUCIONAL DE GESTIÓN DE RIESGOS")
    p_title_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title_bottom.style.font.size = Pt(9)
    p_title_bottom.style.font.name = 'Arial'
    
    # ------------------------------------------------------------------
    # C/D. CONTENIDO: Columnas 3 y 4 (Datos de Control)
    # ------------------------------------------------------------------
    control_data = [
        ("Código:", "LMM_ORG_05"), 
        ("Rev.:", "00"),
        ("Vigencia:", "00/00/2025"),
        ("Página:", "1-1")
    ]

    for i, (key, value) in enumerate(control_data):
        c_cell = header_table.cell(i, 2)
        d_cell = header_table.cell(i, 3)
        
        # Columna 3 (Etiqueta)
        p_key = c_cell.paragraphs[0]
        run_key = p_key.add_run(key)
        run_key.bold = True
        run_key.font.size = Pt(8) 
        run_key.font.name = 'Arial'
        p_key.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Columna 4 (Valor)
        p_value = d_cell.paragraphs[0]
        run_value = p_value.add_run(value)
        run_value.font.size = Pt(8)
        run_value.font.name = 'Arial'
        p_value.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # --------------------------------------------------------
    # 3. PIE DE PÁGINA (AVISO DE CONFIDENCIALIDAD)
    # --------------------------------------------------------
    footer = section.footer
    footer.paragraphs[0].clear() # Limpiar el párrafo predeterminado
    
    # Línea 1
    p_line1 = footer.paragraphs[0] if len(footer.paragraphs) > 0 else footer.add_paragraph()
    p_line1.clear()
    run1 = p_line1.add_run("Este documento contiene información de propiedad exclusiva de La Química Farmacéutica S.A. Queda prohibida la difusión")
    run1.bold = True
    run1.font.size = Pt(8)
    run1.font.name = 'Arial'
    p_line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Línea 2
    p_line2 = footer.add_paragraph()
    run2 = p_line2.add_run("y/o cesión a terceros sin autorización previa del área de Auditoría Interna y O&M. Toda copia no controlada carece de validez.")
    run2.bold = True
    run2.font.size = Pt(8)
    run2.font.name = 'Arial'
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --------------------------------------------------------
    # 4. CUERPO DEL DOCUMENTO
    # --------------------------------------------------------

    # Título de la Compañía y Ficha
    p_id = document.add_paragraph()
    p_id.add_run('Lqf La química farmacéutica').bold = True
    p_id.add_run('\nFICHA DE RIESGO').bold = True
    p_id.style.font.size = Pt(14)
    document.add_paragraph('-' * 80) # Separador visual

    # TÍTULO PRINCIPAL DE LA FICHA
    titulo_doc = document.add_heading(f'Identificación de Riesgo N° {datos_riesgo["Num_Riesgo"]}', level=0)
    titulo_doc.style.font.size = Pt(16)
    
    document.add_paragraph(f'Versión: {datos_riesgo["Version"]} | Última Revisión: {datos_riesgo["Ultima_Revision"]}').style.font.size = Pt(9)
    document.add_paragraph() 

    # 1) - IDENTIFICACIÓN DEL RIESGO
    identificacion_data = {
        'Riesgo Identificado': datos_riesgo['Riesgo_Identificado'],
        'Entorno de Control': datos_riesgo['Entorno_Control'],
        'Origen / Área Responsable': datos_riesgo['Origen_Area'],
        'Proceso o Documento': datos_riesgo['Proceso_Documento']
    }
    agregar_seccion_tabla(document, '1) IDENTIFICACIÓN DEL RIESGO', identificacion_data)

    # 2) - ANÁLISIS DEL RIESGO
    analisis_data = {
        'Impacto Potencial': datos_riesgo['Impacto_Potencial'],
        'Efecto (Consecuencias)': datos_riesgo['Efecto']
    }
    agregar_seccion_tabla(document, '2) ANÁLISIS DEL RIESGO', analisis_data)

    # 3) - EVALUACIÓN DEL RIESGO
    document.add_heading('3) EVALUACIÓN DEL RIESGO', level=2)
    
    # Tabla específica para la evaluación PxG
    tabla_evaluacion = document.add_table(rows=2, cols=4)
    tabla_evaluacion.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = tabla_evaluacion.rows[0].cells
    for i, text in enumerate(['Gravedad (G)', 'Probabilidad (P)', 'Resultado (P x G)', 'ESCALA DE RIESGO']):
        run = hdr_cells[i].paragraphs[0].add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Valores
    val_cells = tabla_evaluacion.rows[1].cells
    
    # Usar .text garantiza Arial por el estilo 'Normal'
    val_cells[0].text = str(datos_riesgo['Gravedad'])
    val_cells[1].text = str(datos_riesgo['Probabilidad'])
    val_cells[2].text = str(datos_riesgo['PxG'])
    
    escala_run = val_cells[3].paragraphs[0].add_run(str(datos_riesgo['Escala_Riesgo']).upper())
    escala_run.bold = True
    escala_run.font.name = 'Arial'
    val_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # 4) - SEGUIMIENTO DEL RIESGO
    seguimiento_data = {
        'Responsable del Seguimiento': datos_riesgo['Responsable_Seguimiento'],
        'Tipo de Control': datos_riesgo['Tipo_Control'],
        'Eficacia del Seguimiento': datos_riesgo['Eficacia_Seguimiento']
    }
    agregar_seccion_tabla(document, '4) SEGUIMIENTO DEL RIESGO', seguimiento_data)
    
    document.add_heading('Descripción del Control Existente', level=3)
    document.add_paragraph(datos_riesgo['Control_Existente'])

    # 5) - SEGUIMIENTO DE VERSIONES Y ACCIONES
    document.add_heading('5) SEGUIMIENTO DE VERSIONES Y ACCIONES', level=2)
    
    tabla_versiones = document.add_table(rows=1, cols=3)
    tabla_versiones.style = 'Table Grid'
    vers_cells = tabla_versiones.rows[0].cells
    
    run_vers = vers_cells[0].paragraphs[0].add_run(f'Versión: {datos_riesgo["Version"]}')
    run_vers.bold = True
    run_vers.font.name = 'Arial'

    run_est = vers_cells[1].paragraphs[0].add_run(f'Estado: {datos_riesgo["Estado_Control"]}')
    run_est.bold = True
    run_est.font.name = 'Arial'
    
    run_fec = vers_cells[2].paragraphs[0].add_run(f'Fecha Ident.: {datos_riesgo["Fecha_Identificacion"]}')
    run_fec.bold = True
    run_fec.font.name = 'Arial'
    document.add_paragraph()
    
    document.add_heading('Acciones Pendientes / Recomendadas', level=3)
    document.add_paragraph(datos_riesgo['Acciones'])
    
    # Guardar en memoria (BytesIO) y devolver
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --- FUNCIÓN DE CARGA DE DATOS (Se mantiene igual) ---
@st.cache_data
def cargar_datos(archivo, hoja, encabezados, columnas):
    """Carga y procesa los datos del Excel (Cache para rendimiento)."""
    try:
        df = pd.read_excel(
            archivo, 
            sheet_name=hoja, 
            header=encabezados,
            usecols="B:U"
        )
        df = df.fillna("")
        df.columns = columnas
        # Filtra filas que no tienen Número de Riesgo
        df = df[df['Num_Riesgo'] != ""].reset_index(drop=True)
        return df
    except FileNotFoundError:
        st.error(f"Error: No se encontró el archivo de matriz '{archivo}'. Asegúrate de que está en la misma carpeta.")
        return pd.DataFrame()
    except Exception as e:
        st.error(f"Error al leer el Excel: {e}")
        return pd.DataFrame()


# --- INTERFAZ STREAMLIT PRINCIPAL (Se mantiene igual) ---
st.set_page_config(page_title="Generador de Fichas de Riesgo", layout="wide")

st.title("Generador de Fichas de Riesgo Individuales")
st.markdown("Por favor, **ingresa el número de identificación** del riesgo que deseas desplegar y descargar (ej: R-001, R-15).")

df_riesgos = cargar_datos(ARCHIVO_EXCEL, NOMBRE_HOJA, FILA_ENCABEZADOS, COLUMNAS_MAP)

if not df_riesgos.empty:
    
    # Usamos st.text_input para pedir el número de riesgo
    num_riesgo_ingresado = st.sidebar.text_input(
        "Ingresa el Número de Riesgo (ej: R-01)",
        key="riesgo_input"
    ).strip() 

    if num_riesgo_ingresado:
        # BÚSQUEDA: Buscar el riesgo en el DataFrame por el número ingresado
        registro_riesgo_encontrado = df_riesgos[df_riesgos['Num_Riesgo'] == num_riesgo_ingresado]
        
        if not registro_riesgo_encontrado.empty:
            
            registro_riesgo = registro_riesgo_encontrado.iloc[0]
            
            st.header(f"Ficha Seleccionada: {registro_riesgo['Riesgo_Identificado']}")
            
            # Botón de Generación y Descarga
            with st.spinner("Generando ficha..."):
                ficha_docx = generar_ficha_docx(registro_riesgo)
                
                st.download_button(
                    label="📥 Descargar Ficha de Riesgo (DOCX)",
                    data=ficha_docx,
                    file_name=f"Ficha_Riesgo_{registro_riesgo['Num_Riesgo']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.info("Presiona el botón de descarga para obtener el documento A4 generado.")
                
        else:
            # Si no encuentra el número de riesgo
            st.warning(f"⚠️ El número de riesgo '{num_riesgo_ingresado}' no fue encontrado en la matriz. Por favor, verifica el número.")
    
    else:
        st.info("Esperando la entrada del Número de Riesgo...")

else:
    st.error("No se pudieron cargar los datos. Verifica que el archivo Excel sea correcto.")
